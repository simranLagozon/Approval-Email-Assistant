"""
Attachment Parser Service
Supports: PDF, DOCX, TXT, and common text formats
"""

import io
import logging

logger = logging.getLogger(__name__)


def parse_attachment_content(raw_bytes: bytes, filename: str, content_type: str) -> str:
    """
    Parse attachment bytes and return extracted text.
    Returns empty string on failure.
    """
    filename_lower = filename.lower()
    content_type_lower = content_type.lower()

    try:
        if filename_lower.endswith(".pdf") or "pdf" in content_type_lower:
            return _parse_pdf(raw_bytes)
        elif filename_lower.endswith(".docx") or "wordprocessingml" in content_type_lower:
            return _parse_docx(raw_bytes)
        elif filename_lower.endswith(".doc"):
            return _parse_doc(raw_bytes)
        elif filename_lower.endswith(".txt") or "text/plain" in content_type_lower:
            return _parse_txt(raw_bytes)
        elif filename_lower.endswith((".xlsx", ".xls")):
            return _parse_excel(raw_bytes)
        elif filename_lower.endswith(".csv"):
            return _parse_csv(raw_bytes)
        else:
            # Try as plain text
            return _parse_txt(raw_bytes)
    except Exception as e:
        logger.warning(f"Failed to parse attachment {filename}: {e}")
        return f"[Could not parse {filename}: {type(e).__name__}]"


def _parse_pdf(raw_bytes: bytes) -> str:
    """Extract text from PDF using PyMuPDF (fitz) or pdfplumber."""
    # Try PyMuPDF first
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=raw_bytes, filetype="pdf")
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text())
        doc.close()
        return "\n".join(text_parts)
    except ImportError:
        pass

    # Fallback: pdfplumber
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(io.BytesIO(raw_bytes)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
        return "\n".join(text_parts)
    except ImportError:
        pass

    # Last resort: pypdf
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(raw_bytes))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages)
    except Exception as e:
        return f"[PDF parsing unavailable: {e}]"


def _parse_docx(raw_bytes: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(raw_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also extract tables
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    table_text.append(" | ".join(cells))

        return "\n".join(paragraphs + table_text)
    except ImportError:
        return "[python-docx not installed - cannot parse DOCX]"
    except Exception as e:
        return f"[DOCX parse error: {e}]"


def _parse_doc(raw_bytes: bytes) -> str:
    """Attempt to read old .doc files."""
    try:
        # Try antiword or textract if available
        import subprocess
        result = subprocess.run(
            ["antiword", "-"],
            input=raw_bytes,
            capture_output=True,
            timeout=10,
        )
        if result.returncode == 0:
            return result.stdout.decode("utf-8", errors="replace")
    except Exception:
        pass
    return "[.doc format - please convert to .docx for full parsing]"


def _parse_txt(raw_bytes: bytes) -> str:
    """Parse plain text file."""
    for encoding in ["utf-8", "latin-1", "cp1252"]:
        try:
            return raw_bytes.decode(encoding)
        except Exception:
            continue
    return raw_bytes.decode("utf-8", errors="replace")


def _parse_excel(raw_bytes: bytes) -> str:
    """Extract text from Excel files."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(raw_bytes), read_only=True, data_only=True)
        parts = []
        for sheet in wb.worksheets:
            parts.append(f"Sheet: {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                row_vals = [str(c) if c is not None else "" for c in row]
                if any(v.strip() for v in row_vals):
                    parts.append(" | ".join(row_vals))
        return "\n".join(parts)
    except ImportError:
        try:
            import xlrd
            wb = xlrd.open_workbook(file_contents=raw_bytes)
            parts = []
            for sheet in wb.sheets():
                for row_idx in range(sheet.nrows):
                    parts.append(" | ".join(str(sheet.cell_value(row_idx, col)) for col in range(sheet.ncols)))
            return "\n".join(parts)
        except Exception as e:
            return f"[Excel parse error: {e}]"
    except Exception as e:
        return f"[Excel parse error: {e}]"


def _parse_csv(raw_bytes: bytes) -> str:
    """Parse CSV."""
    import csv
    text = _parse_txt(raw_bytes)
    return text[:5000]  # Limit CSV output
